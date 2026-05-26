"""Word 导出服务"""
import io
import structlog
from typing import BinaryIO

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.document import Document
from app.models.citation import Citation

logger = structlog.get_logger()


class ExportService:
    """Word 导出服务"""

    def export_document(
        self,
        document: Document,
        citations: list[Citation] | None = None,
        include_citations: bool = True,
    ) -> BinaryIO:
        """
        将文档导出为 Word 格式
        
        Args:
            document: 文档对象
            citations: 引用列表
            include_citations: 是否包含引用列表
        
        Returns:
            Word 文件二进制流
        """
        doc = DocxDocument()

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(12)

        # 标题
        title = doc.add_heading(document.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 内容（简单处理 HTML，实际应使用更复杂的解析）
        content = self._html_to_docx(doc, document.content_html)

        # 引用列表
        if include_citations and citations:
            doc.add_page_break()
            doc.add_heading("引用来源", level=1)

            for i, citation in enumerate(citations, 1):
                p = doc.add_paragraph()
                p.add_run(f"[{i}] {citation.source_name}").bold = True
                doc.add_paragraph(citation.chunk_content)
                doc.add_paragraph()

        # 保存到 BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logger.info(
            "document_exported",
            document_id=str(document.id),
            title=document.title,
        )

        return buffer

    def _html_to_docx(self, doc: DocxDocument, html: str) -> None:
        """
        将 HTML 内容转换为 Word 文档
        
        TODO: 实现完整的 HTML 解析
        当前为简化实现
        """
        # 简单的文本处理
        import re

        # 移除 HTML 标签，保留文本
        text = re.sub(r"<[^>]+>", "", html)
        # 处理空白字符
        text = re.sub(r"\s+", "\n", text)
        # 移除多余空行
        text = re.sub(r"\n{3,}", "\n\n", text)

        # 分割段落并添加
        for paragraph in text.strip().split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    def _parse_html_tag(self, tag: str) -> dict:
        """解析 HTML 标签，返回样式字典"""
        styles = {}

        tag_lower = tag.lower().strip()

        if tag_lower.startswith("h1"):
            styles["heading"] = 1
        elif tag_lower.startswith("h2"):
            styles["heading"] = 2
        elif tag_lower.startswith("h3"):
            styles["heading"] = 3
        elif tag_lower in ["p", "div"]:
            styles["paragraph"] = True
        elif tag_lower in ["strong", "b"]:
            styles["bold"] = True
        elif tag_lower in ["em", "i"]:
            styles["italic"] = True

        return styles
