"""Markdown → DOCX 导出服务

将 AI 回复的 Markdown 文本转换为格式化的 Word 文档。
支持：标题、段落、粗体/斜体、有序/无序列表、代码块、引用、表格、水平线。
"""
import io
import logging
import re
from typing import TextIO

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# 中文字体名称（宋体）
_CN_FONT = "宋体"
# 等宽字体
_CODE_FONT = "Courier New"


def _set_font(run, font_name: str = _CN_FONT, size_pt: int = 12, bold: bool = False,
              italic: bool = False, color: tuple | None = None):
    """统一设置 run 的字体属性"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    # 东亚文字字体（保证中文生效）
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_inline_text(para, text: str):
    """将含 inline 标记（**bold** *italic* `code`）的文本添加到段落。

    支持：
      - **bold**
      - *italic*  （用单个 * 包裹）
      - `code`
      - [link](url)
    不嵌套处理，按出现顺序依次消费。
    """
    # 用正则按 token 切分：bold / italic / code / link / plain
    pattern = re.compile(
        r'(\*\*.+?\*\*)|'      # **bold**
        r'(\*.+?\*)|'           # *italic*  (非贪婪)
        r'(`.+?`)|'             # `code`
        r'(\[.+?\]\(.+?\))|'   # [text](url)
        r'(.+?)',               # plain text
        re.DOTALL
    )
    # 为了简化，采用顺序扫描法
    i = 0
    s = text
    while i < len(s):
        # **bold**
        if s[i:].startswith("**"):
            end = s.find("**", i + 2)
            if end == -1:
                end = len(s)
            run = para.add_run(s[i + 2:end])
            _set_font(run, bold=True)
            i = end + 2
        # *italic* (确保不是 **)
        elif s[i] == '*' and not s[i:].startswith("**"):
            end = s.find("*", i + 1)
            if end == -1:
                end = len(s)
            run = para.add_run(s[i + 1:end])
            _set_font(run, italic=True)
            i = end + 1
        # `code`
        elif s[i] == '`':
            end = s.find('`', i + 1)
            if end == -1:
                end = len(s)
            run = para.add_run(s[i + 1:end])
            _set_font(run, font_name=_CODE_FONT, size_pt=10)
            i = end + 1
        # [link](url)
        elif s[i] == '[':
            m = re.match(r'\[(.+?)\]\((.+?)\)', s[i:])
            if m:
                run = para.add_run(m.group(1))
                # Word 超链接较复杂，这里先着色 + 下划线示意
                _set_font(run, color=(0, 0, 255))
                run.font.underline = True
                i += m.end()
            else:
                run = para.add_run(s[i])
                _set_font(run)
                i += 1
        else:
            # 普通文本，一直取到下一个特殊字符
            j = i
            while j < len(s) and s[j] not in ('*', '`', '['):
                j += 1
            run = para.add_run(s[i:j])
            _set_font(run)
            i = j
    # 兜底：段落内没有 run 时加一个空 run
    if not list(para.runs):
        run = para.add_run(s)
        _set_font(run)


class MarkdownDocxService:
    """将 Markdown 字符串转换为 python-docx Document 对象"""

    def __init__(self, title: str = ""):
        self.doc = DocxDocument()
        self._setup_styles()
        if title:
            h = self.doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _setup_styles(self):
        """设置文档默认字体为宋体"""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = _CN_FONT
        font.size = Pt(12)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), _CN_FONT)

    # ------------------------------------------------------------------ #
    # 公开接口
    # ------------------------------------------------------------------ #
    def convert(self, md: str) -> DocxDocument:
        """解析 Markdown 文本，填充 self.doc"""
        lines = md.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # ---- 空行 → 跳过（段落由连续非空行组成） ----
            if not stripped:
                i += 1
                continue

            # ---- 代码块（``` 包裹） ----
            if stripped.startswith("```"):
                i = self._parse_code_block(lines, i)
                continue

            # ---- 标题 ----
            heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
            if heading_match:
                level = min(len(heading_match.group(1)), 4)
                text = heading_match.group(2).strip()
                h = self.doc.add_heading(text, level=level)
                # 中文字体
                for run in h.runs:
                    _set_font(run, size_pt=16 - (level - 1) * 2)
                i += 1
                continue

            # ---- 水平线 ----
            if re.match(r'^[-*_]{3,}\s*$', stripped):
                self.doc.add_paragraph("─" * 40)
                i += 1
                continue

            # ---- 表格（简单 | 分隔符语法） ----
            if stripped.startswith("|") and i + 1 < len(lines) and re.match(r'^\|?[\s\-:|]+\|?$', lines[i + 1].strip()):
                i = self._parse_table(lines, i)
                continue

            # ---- 无序列表 ----
            if re.match(r'^[\*\-\+]\s+', stripped):
                i = self._parse_list(lines, i, ordered=False)
                continue

            # ---- 有序列表 ----
            if re.match(r'^\d+\.\s+', stripped):
                i = self._parse_list(lines, i, ordered=True)
                continue

            # ---- 引用 ----
            if stripped.startswith(">"):
                i = self._parse_blockquote(lines, i)
                continue

            # ---- 普通段落 ----
            i = self._parse_paragraph(lines, i)

        return self.doc

    def to_bytes(self) -> bytes:
        """将文档序列化为 bytes"""
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ------------------------------------------------------------------ #
    # 内部解析方法
    # ------------------------------------------------------------------ #
    def _parse_code_block(self, lines, start):
        """解析 ``` 代码块，使用等宽字体"""
        # 跳过开始的 ```
        lang = lines[start].strip()[3:].strip()
        i = start + 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i])
            i += 1
        # 添加代码段落
        para = self.doc.add_paragraph()
        para.style = "No Spacing"
        run = para.add_run("\n".join(code_lines))
        _set_font(run, font_name=_CODE_FONT, size_pt=9)
        # 灰色底纹（通过段落边框模拟）
        para_format = para.paragraph_format
        para_format.left_indent = Inches(0.3)
        para_format.right_indent = Inches(0.3)
        i += 1  # 跳过结束的 ```
        return i

    def _parse_table(self, lines, start):
        """解析 Markdown 表格（| 分隔符）"""
        # 第一行：表头
        header = [c.strip() for c in lines[start].strip().strip("|").split("|")]
        # 第二行：对齐规则（跳过）
        i = start + 2
        rows = [header]
        while i < len(lines) and lines[i].strip().startswith("|"):
            row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(row)
            i += 1

        if len(rows) < 1:
            return i

        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"

        for ri, row_data in enumerate(rows):
            for ci, cell_text in enumerate(row_data):
                cell = table.cell(ri, ci)
                cell.text = cell_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        _set_font(run, size_pt=10)
                if ri == 0:  # 表头加粗
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True

        self.doc.add_paragraph()  # 表格后加空行
        return i

    def _parse_list(self, lines, start, ordered: bool):
        """解析有序/无序列表，支持一级嵌套"""
        i = start
        while i < len(lines):
            stripped = lines[i].strip()
            # 判断是否是列表项
            if ordered and re.match(r'^\d+\.\s+', stripped):
                text = re.sub(r'^\d+\.\s+', '', stripped)
                para = self.doc.add_paragraph(style='List Number')
                _add_inline_text(para, text)
                i += 1
            elif not ordered and re.match(r'^[\*\-\+]\s+', stripped):
                text = re.sub(r'^[\*\-\+]\s+', '', stripped)
                para = self.doc.add_paragraph(style='List Bullet')
                _add_inline_text(para, text)
                i += 1
            else:
                break
        return i

    def _parse_blockquote(self, lines, start):
        """解析引用块（> 开头，可多行）"""
        i = start
        quote_lines = []
        while i < len(lines) and (lines[i].strip().startswith(">") or (lines[i].strip() and i == start + 1)):
            quote_lines.append(lines[i].strip().lstrip(">").strip())
            i += 1
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)
        run = para.add_run("\n".join(quote_lines))
        _set_font(run, italic=True)
        return i

    def _parse_paragraph(self, lines, start):
        """解析普通段落（连续非空行直到遇到空行或特殊行）"""
        i = start
        parts = []
        while i < len(lines):
            stripped = lines[i].strip()
            if not stripped:
                break
            # 遇到特殊行就停
            if (stripped.startswith("#") or
                stripped.startswith("```") or
                stripped.startswith(">") or
                stripped.startswith("|") or
                re.match(r'^[\*\-\+]\s+', stripped) or
                re.match(r'^\d+\.\s+', stripped) or
                re.match(r'^[-*_]{3,}\s*$', stripped)):
                break
            parts.append(stripped)
            i += 1
        if parts:
            para = self.doc.add_paragraph()
            _add_inline_text(para, " ".join(parts))
        return i


def markdown_to_docx_bytes(markdown_text: str, title: str = "") -> bytes:
    """一键将 Markdown 文本转换为 .docx 文件 bytes"""
    svc = MarkdownDocxService(title=title)
    svc.convert(markdown_text)
    return svc.to_bytes()
